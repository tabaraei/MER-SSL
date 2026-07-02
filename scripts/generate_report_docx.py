"""
generate_report_docx.py — editable Word (.docx) version of the progress report
==============================================================================
Produces mert_progress_report_v3.docx with the SAME content as the PDF, so it is
editable in Word for minor tweaks. It does NOT re-author the report: it runs the
existing `generate_report_v3.py` with a stubbed reportlab, captures the exact
flowables it builds (same text, same numbers), and renders them to .docx. This
guarantees the Word file never drifts from the audited PDF content.

Run with a Python that has python-docx (conda):
    /opt/miniconda3/bin/python generate_report_docx.py
Output: mert_progress_report_v3.docx
"""

import os
import re

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "generate_report_v3.py")
OUT = os.path.join(BASE, "mert_progress_report_v3.docx")

# ── 1. Stub reportlab so the source builds its flowable list without a PDF ──
class _Style:
    def __init__(self, name="Normal", parent=None, **kw):
        self.name = name; self.kw = kw

class Paragraph:
    def __init__(self, text, style=None): self.text = text; self.style = style
class Spacer:
    def __init__(self, *a, **k): pass
class HRFlowable:
    def __init__(self, *a, **k): pass
class Image:
    def __init__(self, path, *a, **k): self.path = path
class KeepTogether:
    def __init__(self, content, *a, **k): self.content = content
class Table:
    def __init__(self, data, *a, **k): self.data = data
    def setStyle(self, *a, **k): pass
class TableStyle:
    def __init__(self, *a, **k): pass
class SimpleDocTemplate:
    def __init__(self, *a, **k): pass
    def build(self, *a, **k): pass

class _SS(dict):
    def __getitem__(self, k):
        v = dict.get(self, k)
        if v is None:
            v = _Style(k); self[k] = v
        return v
def getSampleStyleSheet(): return _SS()

class _Colors:
    white = "white"
    def HexColor(self, x): return x
colors = _Colors()

A4 = (595.276, 841.89)
cm = 28.3465
TA_LEFT, TA_CENTER, TA_JUSTIFY = 0, 1, 4


def _run_source():
    """Exec generate_report_v3.py with stubs; return its `s` flowable list."""
    src = open(SRC, encoding="utf-8").read()
    # drop reportlab imports: parenthesised multi-line form first, then single-line
    src = re.sub(r"^from reportlab[^\n]*\([^)]*\)\n", "", src, flags=re.MULTILINE)
    src = re.sub(r"^from reportlab.*\n", "", src, flags=re.MULTILINE)
    ns = {
        "os": os, "ParagraphStyle": _Style, "getSampleStyleSheet": getSampleStyleSheet,
        "Paragraph": Paragraph, "Spacer": Spacer, "HRFlowable": HRFlowable,
        "Image": Image, "KeepTogether": KeepTogether, "Table": Table,
        "TableStyle": TableStyle, "SimpleDocTemplate": SimpleDocTemplate,
        "colors": colors, "A4": A4, "cm": cm,
        "TA_LEFT": TA_LEFT, "TA_CENTER": TA_CENTER, "TA_JUSTIFY": TA_JUSTIFY,
        "__name__": "_stub", "__file__": SRC,
    }
    exec(compile(src, SRC, "exec"), ns)
    return ns["s"]


# ── 2. Inline reportlab markup → docx runs ──
_ENT = {"&amp;": "&", "&nbsp;": " ", "&lt;": "<", "&gt;": ">"}

def _emit_runs(paragraph, text):
    """Add runs to a docx paragraph, honouring <b>/<i>/<br/> and entities."""
    for k, v in _ENT.items():
        text = text.replace(k, v)
    # tokenise on tags
    tokens = re.split(r"(<[^>]+>)", text)
    bold = italic = False
    for tok in tokens:
        if not tok:
            continue
        low = tok.lower()
        if low in ("<b>", "<strong>"): bold = True
        elif low in ("</b>", "</strong>"): bold = False
        elif low in ("<i>", "<em>", "<oblique>"): italic = True
        elif low in ("</i>", "</em>", "</oblique>"): italic = False
        elif low in ("<br/>", "<br>", "<br />"): paragraph.add_run().add_break()
        elif tok.startswith("<"):  # ignore other tags (font, etc.)
            continue
        else:
            r = paragraph.add_run(tok); r.bold = bold; r.italic = italic

def _plain(text):
    for k, v in _ENT.items():
        text = text.replace(k, v)
    return re.sub(r"<[^>]+>", "", text).strip()


# ── 3. Render flowables to docx ──
def build_docx(flowables):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    doc = Document()

    NAVY = RGBColor(0x1B, 0x3A, 0x6B)
    TEAL = RGBColor(0x2E, 0x86, 0xAB)

    def render(flow):
        cls = flow.__class__.__name__
        if cls == "KeepTogether":
            content = flow.content if isinstance(flow.content, (list, tuple)) else [flow.content]
            for f in content:
                render(f)
            return
        if cls == "Spacer" or cls == "HRFlowable":
            return
        if cls == "Image":
            p = doc.add_paragraph()
            if os.path.exists(flow.path):
                try:
                    doc.add_picture(flow.path, width=Inches(6.0))
                    return
                except Exception:
                    pass
            r = p.add_run(f"[figure: {os.path.basename(flow.path)}]"); r.italic = True
            return
        if cls == "Table":
            rows = flow.data
            if not rows:
                return
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=len(rows), cols=ncol)
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j in range(ncol):
                    cell_txt = _plain(row[j].text) if j < len(row) and hasattr(row[j], "text") else ""
                    cellp = t.cell(i, j).paragraphs[0]
                    run = cellp.add_run(cell_txt)
                    run.font.size = Pt(8.5)
                    if i == 0:
                        run.bold = True
            doc.add_paragraph()
            return
        if cls == "Paragraph":
            name = getattr(flow.style, "name", "B") if flow.style else "B"
            if name == "T":               # title
                p = doc.add_paragraph(); p.style = doc.styles["Title"]
                _emit_runs(p, flow.text)
                for run in p.runs:
                    run.font.color.rgb = NAVY
            elif name == "S":             # subtitle
                p = doc.add_paragraph(); _emit_runs(p, flow.text)
                for run in p.runs:
                    run.italic = True; run.font.color.rgb = TEAL; run.font.size = Pt(11)
            elif name == "H1":
                doc.add_heading(_plain(flow.text), level=1)
            elif name == "H2":
                doc.add_heading(_plain(flow.text), level=2)
            elif name == "Bu":           # bullet
                p = doc.add_paragraph(style="List Bullet"); _emit_runs(p, flow.text)
            elif name in ("C",):          # caption
                p = doc.add_paragraph(); _emit_runs(p, flow.text)
                for run in p.runs:
                    run.italic = True; run.font.size = Pt(8)
            elif name in ("N",):          # amber note
                p = doc.add_paragraph(); _emit_runs(p, flow.text)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
            else:                          # body / meta
                p = doc.add_paragraph(); _emit_runs(p, flow.text)
            return

    for flow in flowables:
        render(flow)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    flowables = _run_source()
    out = build_docx(flowables)
    n_para = sum(1 for f in flowables if f.__class__.__name__ == "Paragraph")
    n_tab = sum(1 for f in flowables if f.__class__.__name__ == "Table")
    print(f"✅ Word report -> {out}")
    print(f"   captured {len(flowables)} flowables ({n_para} paragraphs, {n_tab} tables)")
